#!/usr/bin/env python3
"""Build the deterministic compact-reference DOCX used by the Tools manual."""

from __future__ import annotations

import argparse
from collections.abc import Sequence
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from scripts.tools_manual_artifacts import canonicalize_docx

BLUE = RGBColor(0x2E, 0x74, 0xB5)
BLUE_DARK = RGBColor(0x1F, 0x4D, 0x78)
BODY = RGBColor(0x20, 0x25, 0x2B)
FONT = "Arial"


def _font(style: object, *, points: float, color: RGBColor, bold: bool) -> None:
    font = style.font  # type: ignore[attr-defined]
    font.name = FONT
    font.size = Pt(points)
    font.color.rgb = color
    font.bold = bold
    properties = style.element.get_or_add_rPr()  # type: ignore[attr-defined]
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.append(fonts)
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)


def _paragraph(style: object, *, before: float, after: float, line: float) -> None:
    paragraph = style.paragraph_format  # type: ignore[attr-defined]
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line
    paragraph.keep_with_next = before > 0


def build_reference(path: Path) -> None:
    """Create the renderer-owned reference DOCX from explicit style tokens."""
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    tokens = {
        "Normal": (11, BODY, False, 0, 6, 1.25),
        "Title": (24, BLUE_DARK, True, 0, 12, 1.0),
        "Subtitle": (13, BODY, False, 0, 14, 1.0),
        "Heading 1": (16, BLUE, True, 18, 10, 1.0),
        "Heading 2": (13, BLUE, True, 14, 7, 1.0),
        "Heading 3": (12, BLUE_DARK, True, 10, 5, 1.0),
        "Caption": (9, BODY, False, 4, 6, 1.0),
    }
    for name, (size, color, bold, before, after, line) in tokens.items():
        style = document.styles[name]
        _font(style, points=size, color=color, bold=bold)
        _paragraph(style, before=before, after=after, line=line)

    header = section.header.paragraphs[0]
    header.text = "Tools Engineering Design Manual | Generated - Unapproved"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x5B, 0x65, 0x73)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    properties = document.core_properties
    properties.title = "Tools Engineering Design Manual reference style"
    properties.author = "D-sorganization"
    properties.last_modified_by = "D-sorganization deterministic renderer"
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    canonicalize_docx(path)


def main(argv: Sequence[str] | None = None) -> int:
    """Build the default reference path or a caller-selected output."""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args(argv)
    root = Path(__file__).resolve().parents[1]
    output = (
        args.output or root / "manuals" / "tools" / "styles" / "tools-reference.docx"
    )
    build_reference(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
